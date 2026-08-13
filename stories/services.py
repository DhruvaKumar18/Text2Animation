import os
import re
import json
import fitz  # PyMuPDF
import pdfplumber
import docx
from django.core.exceptions import ValidationError

class DocumentParserService:
    @classmethod
    def parse_document(cls, file_obj, filename: str) -> str:
        """
        Dispatches file parsing based on extension and returns cleaned extracted text.
        """
        ext = os.path.splitext(filename)[1].lower()
        
        # Reset file stream position to start just in case
        file_obj.seek(0)
        
        if ext == '.txt':
            raw_text = cls.extract_txt(file_obj)
        elif ext == '.pdf':
            raw_text = cls.extract_pdf(file_obj)
        elif ext == '.docx':
            raw_text = cls.extract_docx(file_obj)
        else:
            raise ValidationError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")
            
        cleaned_text = cls.clean_text(raw_text)
        return cleaned_text

    @classmethod
    def extract_txt(cls, file_obj) -> str:
        """
        Extracts text from plain text (.txt) files.
        Tries UTF-8 first, with fallbacks to other common decodings.
        """
        content = file_obj.read()
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValidationError("Could not decode text file with standard encodings (UTF-8, Latin-1).")

    @classmethod
    def extract_pdf(cls, file_obj) -> str:
        """
        Extracts text from PDF files using PyMuPDF (fitz) with a fallback to pdfplumber.
        """
        text_content = []
        
        # 1. Try PyMuPDF (extremely fast and robust)
        try:
            # Read bytes of file
            pdf_bytes = file_obj.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page in doc:
                text_content.append(page.get_text())
            doc.close()
            extracted = "\n".join(text_content).strip()
            if extracted:
                return extracted
        except Exception as e:
            # Log warning or pass to fallback
            pass
            
        # 2. Fallback to pdfplumber (very reliable for complex PDF layouts)
        text_content = []
        try:
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            extracted = "\n".join(text_content).strip()
            if extracted:
                return extracted
        except Exception as e:
            raise ValidationError(f"Failed to parse PDF document: {str(e)}")
            
        raise ValidationError("The uploaded PDF file is empty or text could not be extracted.")

    @classmethod
    def extract_docx(cls, file_obj) -> str:
        """
        Extracts text from DOCX files using python-docx.
        Extracts from paragraphs and tables.
        """
        try:
            doc = docx.Document(file_obj)
            text_content = []
            
            # Read paragraphs
            for paragraph in doc.paragraphs:
                p_text = paragraph.text.strip()
                if p_text:
                    text_content.append(p_text)
                    
            # Read tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # De-duplicate adjacent identical cells (common in merged cells)
                    cleaned_row_cells = []
                    for cell in row_cells_text:
                        if not cleaned_row_cells or cell != cleaned_row_cells[-1]:
                            cleaned_row_cells.append(cell)
                    if cleaned_row_cells:
                        text_content.append(" | ".join(cleaned_row_cells))
                        
            return "\n".join(text_content)
        except Exception as e:
            raise ValidationError(f"Failed to parse DOCX document: {str(e)}")

    @classmethod
    def clean_text(cls, text: str) -> str:
        """
        Cleans extracted text:
        1. Normalizes spaces and tabs within lines.
        2. Discards lines matching common page numbers patterns (e.g. Page 1, Page 1 of 4, 1 / 10).
        3. Collapses consecutive empty lines down to a single empty line.
        """
        if not text:
            return ""
            
        # Regex patterns for page number headers/footers (matching entire line)
        page_patterns = [
            re.compile(r'^\s*page\s*\d+\s*$', re.IGNORECASE),
            re.compile(r'^\s*page\s*\d+\s*of\s*\d+\s*$', re.IGNORECASE),
            re.compile(r'^\s*\d+\s*/\s*\d+\s*$', re.IGNORECASE),
            re.compile(r'^\s*\[?\s*\d+\s*\]?\s*$', re.IGNORECASE), # stand-alone numbers, e.g. "1" or "[1]"
        ]
        
        lines = text.splitlines()
        cleaned_lines = []
        
        for line in lines:
            # Normalize whitespace within the line (replace tabs and multiple spaces with a single space)
            line_str = re.sub(r'[ \t]+', ' ', line).strip()
            
            # Skip if it matches any page number pattern
            is_page_number = False
            for pattern in page_patterns:
                if pattern.match(line_str):
                    is_page_number = True
                    break
                    
            if is_page_number:
                continue
                
            cleaned_lines.append(line_str)
            
        # Collapse multiple consecutive empty lines
        final_lines = []
        for line in cleaned_lines:
            if not line:
                # If current line is empty, only add it if the previous line wasn't empty
                if final_lines and final_lines[-1] != "":
                    final_lines.append("")
            else:
                final_lines.append(line)
                
        # Strip leading/trailing empty lines
        while final_lines and final_lines[0] == "":
            final_lines.pop(0)
        while final_lines and final_lines[-1] == "":
            final_lines.pop()
            
        return "\n".join(final_lines)


class ImageGeneratorService:
    @classmethod
    def generate_image_for_scene(cls, scene, output_path: str) -> dict:
        """
        Orchestrates the fallback image generation chain:
        HuggingFace -> Fal.ai -> Replicate -> Pillow Mock
        Applies a retry mechanism (1 retry per provider) on failure,
        logs durations and seeds, and saves visual assets to output_path.
        """
        import time
        import random
        from django.conf import settings
        
        # Ensure directory path exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        prompt = scene.image_prompt or scene.prompt or "Cinematic visual landscape"
        
        # Determine seed
        if scene.image_seed:
            seed = scene.image_seed
        else:
            seed = random.randint(1, 999999999999)
            scene.image_seed = seed
            scene.save()
            
        hf_key = getattr(settings, 'HF_API_KEY', '') or os.environ.get('HF_API_KEY', '')
        fal_key = getattr(settings, 'FAL_KEY', '') or os.environ.get('FAL_KEY', '')
        replicate_key = getattr(settings, 'REPLICATE_API_TOKEN', '') or os.environ.get('REPLICATE_API_TOKEN', '')
        
        providers = []
        if hf_key:
            providers.append(('HUGGINGFACE', 'FLUX.1-schnell', lambda p, s: cls._call_huggingface(p, s, hf_key)))
        if fal_key:
            providers.append(('FAL_AI', 'fal-ai/flux/schnell', lambda p, s: cls._call_fal(p, s, fal_key)))
        if replicate_key:
            providers.append(('REPLICATE', 'flux-schnell', lambda p, s: cls._call_replicate(p, s, replicate_key)))
            
        if not providers:
            providers.append(('MOCK', 'MOCK_PILLOW', lambda p, s: None))
            
        image_bytes = None
        model_used = None
        provider_used = None
        start_time = time.time()
        
        for provider, model_name, api_call in providers:
            if provider == 'MOCK':
                break
                
            success = False
            # Try once + 1 retry = 2 attempts
            for attempt in range(2):
                try:
                    image_bytes = api_call(prompt, seed)
                    if image_bytes:
                        model_used = model_name
                        provider_used = provider
                        success = True
                        break
                except Exception as e:
                    # Let it log or continue
                    time.sleep(0.5) # Brief delay before retry
                    
            if success:
                break
                
        duration = time.time() - start_time
        
        if image_bytes:
            # Save the retrieved binary content
            with open(output_path, 'wb') as f:
                f.write(image_bytes)
        else:
            # Fallback to local Pillow image builder
            cls._local_pillow_mock_image(prompt, scene.scene_number, output_path)
            provider_used = 'MOCK'
            model_used = 'MOCK_PILLOW'
            
        # Update scene record with metadata
        scene.image_model = model_used
        scene.image_seed = seed
        scene.image_generation_time = duration
        scene.save()
        
        return {
            'provider': provider_used,
            'model': model_used,
            'seed': seed,
            'duration': duration
        }

    @classmethod
    def _call_huggingface(cls, prompt: str, seed: int, api_key: str) -> bytes:
        import requests
        url = "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-schnell"
        headers = {
            'Authorization': f"Bearer {api_key}",
            'Content-Type': 'application/json'
        }
        payload = {
            "inputs": prompt,
            "parameters": {
                "seed": seed
            }
        }
        response = requests.post(url, headers=headers, json=payload, timeout=12)
        response.raise_for_status()
        return response.content

    @classmethod
    def _call_fal(cls, prompt: str, seed: int, api_key: str) -> bytes:
        import requests
        url = "https://fal.run/fal-ai/flux/schnell"
        headers = {
            'Authorization': f"Key {api_key}",
            'Content-Type': 'application/json'
        }
        payload = {
            "prompt": prompt,
            "seed": seed,
            "sync_mode": True
        }
        response = requests.post(url, headers=headers, json=payload, timeout=12)
        response.raise_for_status()
        data = response.json()
        img_url = data['images'][0]['url']
        
        # Download the actual image bytes
        img_res = requests.get(img_url, timeout=8)
        img_res.raise_for_status()
        return img_res.content

    @classmethod
    def _call_replicate(cls, prompt: str, seed: int, api_key: str) -> bytes:
        import requests
        import time
        url = "https://api.replicate.com/v1/predictions"
        headers = {
            'Authorization': f"Bearer {api_key}",
            'Content-Type': 'application/json'
        }
        payload = {
            "version": "0e1def17a16dbf111b5021f7c82353bcac955b3cc0a7aa05468d37fd55b79fc9",
            "input": {
                "prompt": prompt,
                "seed": seed
            }
        }
        # Start prediction
        response = requests.post(url, headers=headers, json=payload, timeout=10)
        response.raise_for_status()
        prediction = response.json()
        get_url = prediction['urls']['get']
        
        # Poll status
        for _ in range(8):
            time.sleep(1.5)
            poll_res = requests.get(get_url, headers=headers, timeout=5)
            poll_res.raise_for_status()
            pred_data = poll_res.json()
            if pred_data['status'] == 'succeeded':
                output_url = pred_data['output'][0]
                img_res = requests.get(output_url, timeout=8)
                img_res.raise_for_status()
                return img_res.content
            elif pred_data['status'] in ['failed', 'canceled']:
                raise ValueError(f"Replicate prediction failed: {pred_data.get('error')}")
                
        raise TimeoutError("Replicate prediction timed out.")

    @classmethod
    def _local_pillow_mock_image(cls, prompt: str, scene_number: int, output_path: str):
        from PIL import Image, ImageDraw, ImageFont
        import random
        width, height = 1280, 720
        img = Image.new('RGB', (width, height), color=(20, 24, 33))
        draw = ImageDraw.Draw(img)
        
        # Gradient backdrop
        r_start, g_start, b_start = random.randint(20, 80), random.randint(20, 80), random.randint(100, 200)
        r_end, g_end, b_end = random.randint(100, 200), random.randint(20, 80), random.randint(120, 220)
        for y in range(height):
            ratio = y / height
            r = int(r_start * (1 - ratio) + r_end * ratio)
            g = int(g_start * (1 - ratio) + g_end * ratio)
            b = int(b_start * (1 - ratio) + b_end * ratio)
            draw.line([(0, y), (width, y)], fill=(r, g, b))
            
        # Draw some visuals
        draw.ellipse([width//4, height//4, width*3//4, height*3//4], outline=(255,255,255,40), width=3)
        draw.polygon([(width//2, height//3), (width//3, height*2//3), (width*2//3, height*2//3)], outline=(255,255,255,20), width=2)
        draw.rectangle([40, 40, width - 40, height - 40], outline=(255, 255, 255, 50), width=2)
        
        try:
            font_title = ImageFont.truetype("arial.ttf", 40)
            font_sub = ImageFont.truetype("arial.ttf", 20)
        except IOError:
            font_title = ImageFont.load_default()
            font_sub = ImageFont.load_default()
            
        draw.text((80, 80), f"SCENE {scene_number}", fill=(255, 255, 255), font=font_title)
        
        # Wrapped prompt text
        words = prompt.split()
        lines = []
        current_line = []
        for word in words:
            if len(" ".join(current_line + [word])) * 8 < width - 200:
                current_line.append(word)
            else:
                lines.append(" ".join(current_line))
                current_line = [word]
        if current_line:
            lines.append(" ".join(current_line))
            
        y_text = 160
        for line in lines[:4]:
            draw.text((80, y_text), line, fill=(240, 240, 255), font=font_sub)
            y_text += 30
            
        img.save(output_path, "JPEG")



class StoryPolisherService:
    POLISH_PROMPT_TEMPLATE = (
        "You are an expert storyteller and script editor. "
        "Please polish, correct, and improve the storytelling of the following story prompt. "
        "Your instructions:\n"
        "1. Correct any spelling, grammar, and punctuation mistakes.\n"
        "2. Enhance the descriptive language, narrative flow, and imagery (make it more cinematic and engaging).\n"
        "3. Do NOT change the core meaning, main plot, characters, or actions.\n"
        "4. Keep the output clean. Return ONLY the polished story text. Do not include any introductory or concluding comments, "
        "explanations, or markdown backticks around the text.\n\n"
        "Original story:\n"
        "{original_text}"
    )

    @classmethod
    def polish_story_text(cls, original_text: str, story_id: int = None) -> dict:
        """
        Orchestrates the fallback polishing chain:
        Google Gemini -> Groq -> OpenRouter -> Mock Fallback
        Measures execution speed, writes logs to the DB, and returns the polished text details.
        """
        import time
        from django.conf import settings
        from history.models import LlmApiLog
        
        prompt = cls.POLISH_PROMPT_TEMPLATE.format(original_text=original_text)
        
        # Fallback chain configurations
        gemini_key = getattr(settings, 'GEMINI_API_KEY', '') or os.environ.get('GEMINI_API_KEY', '')
        groq_key = getattr(settings, 'GROQ_API_KEY', '') or os.environ.get('GROQ_API_KEY', '')
        openrouter_key = getattr(settings, 'OPENROUTER_API_KEY', '') or os.environ.get('OPENROUTER_API_KEY', '')
        
        # Track pipeline attempts
        providers_to_try = []
        if gemini_key:
            providers_to_try.append(('GEMINI', lambda t: cls._call_gemini(t, gemini_key)))
        if groq_key:
            providers_to_try.append(('GROQ', lambda t: cls._call_groq(t, groq_key)))
        if openrouter_key:
            providers_to_try.append(('OPENROUTER', lambda t: cls._call_openrouter(t, openrouter_key)))
            
        # If no real keys are supplied, fallback to Mock directly
        if not providers_to_try:
            providers_to_try.append(('MOCK', lambda t: cls._local_mock_polisher(original_text)))
            
        final_text = None
        provider_used = None
        error_logs = []
        
        start_time = time.time()
        
        for provider_name, api_call in providers_to_try:
            step_start = time.time()
            try:
                final_text = api_call(prompt)
                provider_used = provider_name
                step_duration = time.time() - step_start
                
                # Log success to DB
                LlmApiLog.objects.create(
                    story_id=story_id,
                    provider=provider_name,
                    prompt=prompt,
                    response_text=final_text,
                    status=LlmApiLog.Status.SUCCESS,
                    execution_time_seconds=step_duration
                )
                break  # Successful call, exit loop
            except Exception as e:
                step_duration = time.time() - step_start
                err_msg = str(e)
                error_logs.append(f"[{provider_name}] Failed: {err_msg}")
                
                # Log failure to DB
                LlmApiLog.objects.create(
                    story_id=story_id,
                    provider=provider_name,
                    prompt=prompt,
                    status=LlmApiLog.Status.FAILED,
                    error_message=err_msg,
                    execution_time_seconds=step_duration
                )
                continue
                
        # Ultimate fallback if all configured real APIs failed
        if not final_text:
            step_start = time.time()
            final_text = cls._local_mock_polisher(original_text)
            provider_used = 'MOCK'
            step_duration = time.time() - step_start
            
            LlmApiLog.objects.create(
                story_id=story_id,
                provider='MOCK',
                prompt=prompt,
                response_text=final_text,
                status=LlmApiLog.Status.SUCCESS,
                error_message="All external LLM providers failed. Fallback to mock: " + "; ".join(error_logs),
                execution_time_seconds=step_duration
            )

        total_duration = time.time() - start_time
        return {
            'polished_text': final_text,
            'provider': provider_used,
            'execution_time': total_duration
        }

    @classmethod
    def _call_gemini(cls, prompt: str, api_key: str) -> str:
        import requests
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
        headers = {'Content-Type': 'application/json'}
        payload = {
            "contents": [{
                "parts": [{"text": prompt}]
            }]
        }
        response = requests.post(url, headers=headers, json=payload, timeout=5)
        response.raise_for_status()
        data = response.json()
        try:
            return data['candidates'][0]['content']['parts'][0]['text'].strip()
        except (KeyError, IndexOffBoundsError, IndexError):
            raise ValueError("Malformed response received from Google Gemini API.")

    @classmethod
    def _call_groq(cls, prompt: str, api_key: str) -> str:
        import requests
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            'Authorization': f"Bearer {api_key}",
            'Content-Type': 'application/json'
        }
        payload = {
            "model": "llama3-8b-8192",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2
        }
        response = requests.post(url, headers=headers, json=payload, timeout=5)
        response.raise_for_status()
        data = response.json()
        return data['choices'][0]['message']['content'].strip()

    @classmethod
    def _call_openrouter(cls, prompt: str, api_key: str) -> str:
        import requests
        url = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            'Authorization': f"Bearer {api_key}",
            'Content-Type': 'application/json',
            'HTTP-Referer': 'http://127.0.0.1:8000',
            'X-Title': 'Text2Animation Studio'
        }
        payload = {
            "model": "meta-llama/llama-3-8b-instruct:free",
            "messages": [{"role": "user", "content": prompt}]
        }
        response = requests.post(url, headers=headers, json=payload, timeout=5)
        response.raise_for_status()
        data = response.json()
        return data['choices'][0]['message']['content'].strip()

    @classmethod
    def _local_mock_polisher(cls, original_text: str) -> str:
        """
        Mock text polisher fallback when no keys are supplied or all requests fail.
        Performs light grammar corrections and cinematic styling.
        """
        # Collapse multiple spaces
        text = re.sub(r'[ \t]+', ' ', original_text).strip()
        
        # Simple word enhancements (case insensitive checks)
        enhancements = {
            r'\bastronaut\b': 'cinematic space explorer',
            r'\brobots\b': 'sentient automaton units',
            r'\bcar\b': 'sleek hovercraft vehicle',
            r'\bcity\b': 'cyberpunk metropolis skyline',
            r'\blandscape\b': 'breathtaking visual vistas',
        }
        
        polished = text
        for pattern, replacement in enhancements.items():
            polished = re.sub(pattern, replacement, polished, flags=re.IGNORECASE)
            
        # Ensure it has a capital start and finishes with a period if not present
        if polished:
            polished = polished[0].upper() + polished[1:]
            if not polished.endswith(('.', '!', '?')):
                polished += '.'
                
        # Append a generic polishing signoff decoration
        if polished == text:
            polished = f"[AI Enhanced] A detailed scenario of: {text}"
            
        return polished


class SceneSplitterService:
    SPLIT_PROMPT_TEMPLATE = (
        "You are an expert cinematic director and storyboard artist.\n"
        "Please divide the following story script into a sequence of distinct cinematic scenes.\n"
        "For each scene, extract and compile detailed visual attributes.\n"
        "Your output MUST be a valid JSON array, containing objects with exactly the following key-value structure:\n"
        "[\n"
        "  {{\n"
        "    \"scene_number\": 1,\n"
        "    \"title\": \"Scene Title\",\n"
        "    \"description\": \"Action and description of what occurs in the scene\",\n"
        "    \"characters\": \"Characters involved or none\",\n"
        "    \"environment\": \"Setting location details\",\n"
        "    \"lighting\": \"Lighting characteristics\",\n"
        "    \"mood\": \"Emotional mood\",\n"
        "    \"camera_angle\": \"Camera shot type and angle description\",\n"
        "    \"image_prompt\": \"A detailed visual prompt for generating a high-quality still image (e.g. photorealistic, detailed style, etc.)\",\n"
        "    \"animation_prompt\": \"Prompt describing motion or camera moves (e.g. slow pan, zoom, drift)\",\n"
        "    \"narration\": \"Speech narration or voiceover overlay for this scene\",\n"
        "    \"duration\": 4.0\n"
        "  }}\n"
        "]\n\n"
        "Ensure the output contains ONLY the raw JSON array code block. Do not add conversational text, markdown wrapping (other than json block), or header text.\n\n"
        "Story Script:\n"
        "{story_text}"
    )

    @classmethod
    def split_story_into_scenes(cls, story_text: str, story_id: int = None) -> list[dict]:
        """
        Orchestrates the fallback scene splitting chain:
        Google Gemini -> Groq -> OpenRouter -> Mock Fallback
        Measures execution speed, writes logs to the DB, parses LLM JSON outputs, and returns the scene array.
        """
        import time
        from django.conf import settings
        from history.models import LlmApiLog
        
        prompt = cls.SPLIT_PROMPT_TEMPLATE.format(story_text=story_text)
        
        # Fallback chain configurations
        gemini_key = getattr(settings, 'GEMINI_API_KEY', '') or os.environ.get('GEMINI_API_KEY', '')
        groq_key = getattr(settings, 'GROQ_API_KEY', '') or os.environ.get('GROQ_API_KEY', '')
        openrouter_key = getattr(settings, 'OPENROUTER_API_KEY', '') or os.environ.get('OPENROUTER_API_KEY', '')
        
        providers_to_try = []
        if gemini_key:
            providers_to_try.append(('GEMINI', lambda t: cls._call_gemini(t, gemini_key)))
        if groq_key:
            providers_to_try.append(('GROQ', lambda t: cls._call_groq(t, groq_key)))
        if openrouter_key:
            providers_to_try.append(('OPENROUTER', lambda t: cls._call_openrouter(t, openrouter_key)))
            
        if not providers_to_try:
            providers_to_try.append(('MOCK', lambda t: cls._local_mock_splitter(story_text)))
            
        final_scenes = None
        provider_used = None
        error_logs = []
        
        start_time = time.time()
        
        for provider_name, api_call in providers_to_try:
            step_start = time.time()
            try:
                response_text = api_call(prompt)
                provider_used = provider_name
                step_duration = time.time() - step_start
                
                # Try parsing JSON response
                parsed_json = cls._clean_and_parse_json(response_text)
                if parsed_json and isinstance(parsed_json, list):
                    final_scenes = parsed_json
                    
                    # Log success to DB
                    LlmApiLog.objects.create(
                        story_id=story_id,
                        provider=provider_name,
                        prompt=prompt,
                        response_text=response_text,
                        status=LlmApiLog.Status.SUCCESS,
                        execution_time_seconds=step_duration
                    )
                    break
                else:
                    raise ValueError("Parsed output was not a JSON list array structure.")
            except Exception as e:
                step_duration = time.time() - step_start
                err_msg = str(e)
                error_logs.append(f"[{provider_name}] Failed: {err_msg}")
                
                # Log failure to DB
                LlmApiLog.objects.create(
                    story_id=story_id,
                    provider=provider_name,
                    prompt=prompt,
                    status=LlmApiLog.Status.FAILED,
                    error_message=err_msg,
                    execution_time_seconds=step_duration
                )
                continue
                
        # Ultimate fallback
        if not final_scenes:
            step_start = time.time()
            final_scenes = cls._local_mock_splitter(story_text)
            provider_used = 'MOCK'
            step_duration = time.time() - step_start
            
            LlmApiLog.objects.create(
                story_id=story_id,
                provider='MOCK',
                prompt=prompt,
                response_text=json.dumps(final_scenes),
                status=LlmApiLog.Status.SUCCESS,
                error_message="All LLM splitters failed or returned invalid JSON. Fallback to mock: " + "; ".join(error_logs),
                execution_time_seconds=step_duration
            )

        return final_scenes

    @classmethod
    def _clean_and_parse_json(cls, text: str) -> list:
        # Strip markdown json backticks if present
        cleaned = text.strip()
        match = re.search(r'```(?:json)?\s*(.*?)\s*```', cleaned, re.DOTALL | re.IGNORECASE)
        if match:
            cleaned = match.group(1).strip()
            
        # Try raw json loads
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            # Try to find content starting with [ and ending with ]
            match_array = re.search(r'(\[.*\])', cleaned, re.DOTALL)
            if match_array:
                try:
                    return json.loads(match_array.group(1).strip())
                except json.JSONDecodeError:
                    pass
        return None

    @classmethod
    def _call_gemini(cls, prompt: str, api_key: str) -> str:
        import requests
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
        headers = {'Content-Type': 'application/json'}
        payload = {
            "contents": [{
                "parts": [{"text": prompt}]
            }]
        }
        response = requests.post(url, headers=headers, json=payload, timeout=8)
        response.raise_for_status()
        data = response.json()
        try:
            return data['candidates'][0]['content']['parts'][0]['text'].strip()
        except (KeyError, IndexError):
            raise ValueError("Malformed response received from Google Gemini API.")

    @classmethod
    def _call_groq(cls, prompt: str, api_key: str) -> str:
        import requests
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            'Authorization': f"Bearer {api_key}",
            'Content-Type': 'application/json'
        }
        payload = {
            "model": "llama-3.1-8b-instant",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1
        }
        response = requests.post(url, headers=headers, json=payload, timeout=8)
        response.raise_for_status()
        data = response.json()
        return data['choices'][0]['message']['content'].strip()

    @classmethod
    def _call_openrouter(cls, prompt: str, api_key: str) -> str:
        import requests
        url = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            'Authorization': f"Bearer {api_key}",
            'Content-Type': 'application/json',
            'HTTP-Referer': 'http://127.0.0.1:8000',
            'X-Title': 'Text2Animation Studio'
        }
        payload = {
            "model": "meta-llama/llama-3-8b-instruct:free",
            "messages": [{"role": "user", "content": prompt}]
        }
        response = requests.post(url, headers=headers, json=payload, timeout=8)
        response.raise_for_status()
        data = response.json()
        return data['choices'][0]['message']['content'].strip()

    @classmethod
    def _local_mock_splitter(cls, text: str) -> list[dict]:
        """
        Fallback sentence/paragraph parser to split raw scripts into valid scene arrays.
        """
        # Split text into paragraphs
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        if not paragraphs:
            paragraphs = ["The story begins in a dark, empty space.", "Suddenly, an explosion of light manifests.", "Calm and peace settle across the digital horizon."]
            
        scenes = []
        for idx, paragraph in enumerate(paragraphs, start=1):
            # Extract simple title
            words = paragraph.split()
            title = " ".join(words[:3]).title() if words else f"Scene {idx}"
            
            scenes.append({
                "scene_number": idx,
                "title": title,
                "description": paragraph,
                "characters": "None" if "astronaut" not in paragraph.lower() else "Cinematic space explorer",
                "environment": "Deep space canvas" if "space" in paragraph.lower() or "cosmic" in paragraph.lower() else "Cinematic digital background",
                "lighting": "Volumetric lighting" if "light" in paragraph.lower() else "Ambient glow",
                "mood": "Eerie" if "mysterious" in paragraph.lower() else "Optimistic",
                "camera_angle": "Wide angle pan" if idx % 2 == 0 else "Establishing macro shot",
                "image_prompt": f"Widescreen masterpiece, {paragraph}, cinematic color grading, high detail, photorealistic",
                "animation_prompt": "Slow pan forward, gentle particles drift, 4k",
                "narration": paragraph,
                "duration": float((len(words) // 4) + 3.0)
            })
        return scenes


